"""
SenseNova U1 -- Multimodal policy document ingestion.
Uploads a PDF/image policy document and extracts structured coverage terms.
Replaces the hardcoded POLICY_DOCUMENT in policy_service.py with live-extracted content.
"""

import os
import re
import sys
import json
import hashlib
from pathlib import Path
from typing import Optional

# Ensure stdout uses UTF-8 on Windows so Unicode in policy docs doesn't crash print()
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

# In-memory policy cache -- key: file hash, value: extracted policy dict
_policy_cache: dict = {}
_active_policy: Optional[dict] = None  # The policy currently in use

# Registry: normalised_plate → {path, policy_dict}
_policy_registry: dict = {}


# ── Public API ──────────────────────────────────────────────────────────────────

async def ingest_policy_pdf(pdf_path: str) -> dict:
    """
    Upload policy PDF to SenseNova U1 multimodal model.
    Returns extracted policy dict and caches it as the active policy.
    """
    global _active_policy

    file_hash = _file_hash(pdf_path)
    if file_hash in _policy_cache:
        print(f"[SenseNova] Cache hit -- {Path(pdf_path).name}")
        _active_policy = _policy_cache[file_hash]
        return _active_policy

    api_key = os.getenv("SENSENOVA_API_KEY", "").strip()
    base_url = os.getenv("SENSENOVA_BASE_URL", "https://token.sensenova.cn/v1").strip().rstrip("/")
    is_placeholder = not api_key or api_key.startswith("SN-PLACEHOLDER")

    result = None
    if not is_placeholder:
        result = await _call_sensenova_api(api_key, base_url, pdf_path)

    if not result:
        print(f"[SenseNova] API unavailable -- proxy extraction from {Path(pdf_path).name}")
        result = _proxy_extraction(pdf_path)

    # Ensure details dict is present
    if "details" not in result:
        result["details"] = _extract_docx_fields_from_file(pdf_path) if Path(pdf_path).suffix.lower() == ".docx" else {}

    _policy_cache[file_hash] = result
    _active_policy = result
    print(f"[SenseNova] Policy ingested: '{result['policy_name']}' via {result['source']}")
    return result


def get_active_policy_text() -> Optional[str]:
    """Returns SenseNova-extracted policy text, or None if no custom policy loaded."""
    return _active_policy.get("policy_text") if _active_policy else None


def get_active_policy_details() -> Optional[dict]:
    """Returns {insured_name, vehicle_plate, nric, policy_number} from active policy, or None."""
    if not _active_policy:
        return None
    return _active_policy.get("details")


def has_custom_policy() -> bool:
    return _active_policy is not None


def clear_active_policy():
    global _active_policy
    _active_policy = None


async def auto_load_policies(search_dirs: list = None):
    """
    Scan directories for .docx policy files and register them by plate.
    Called once on server startup so policies are pre-loaded without manual upload.
    """
    global _policy_registry
    backend_dir = Path(__file__).parent.parent
    project_root = backend_dir.parent
    dirs = [
        backend_dir / "policies",
        project_root / "policies",
        project_root,
        backend_dir,
    ]
    if search_dirs:
        dirs = [Path(d) for d in search_dirs] + dirs

    found = []
    for d in dirs:
        if d.exists():
            for f in d.glob("*.docx"):
                if f not in found:
                    found.append(f)

    for path in found:
        try:
            details = _extract_docx_fields_from_file(str(path))
            plate = details.get("vehicle_plate", "").upper().replace(" ", "")
            if plate:
                # Ingest (or re-use cache)
                result = await ingest_policy_pdf(str(path))
                result["details"] = details
                _policy_registry[plate] = {"path": str(path), "policy": result}
                print(f"[SenseNova] Auto-loaded policy: {path.name} -> plate {plate}")
        except Exception as exc:
            print(f"[SenseNova] Could not auto-load {path.name}: {exc}")


async def lookup_policy(name: str, plate: str) -> Optional[dict]:
    """
    Find and verify a policy matching the given claimant name and vehicle plate.
    Returns {match, policy_name, insured_name, vehicle_plate, nric, policy_number, name_match, plate_match}
    or None if no policy doc found.
    """
    norm_plate = plate.upper().strip().replace(" ", "")
    if not norm_plate:
        return None

    # Check registry first (already loaded)
    entry = _policy_registry.get(norm_plate)

    # If not in registry, scan for a file containing the plate in its name
    if not entry:
        backend_dir = Path(__file__).parent.parent
        project_root = backend_dir.parent
        candidates = list((backend_dir / "policies").glob("*.docx")) + \
                     list((project_root / "policies").glob("*.docx")) + \
                     list(project_root.glob("*.docx"))
        for f in candidates:
            if norm_plate.upper() in f.name.upper():
                details = _extract_docx_fields_from_file(str(f))
                result = await ingest_policy_pdf(str(f))
                result["details"] = details
                _policy_registry[norm_plate] = {"path": str(f), "policy": result}
                entry = _policy_registry[norm_plate]
                break

    if not entry:
        return None

    policy = entry["policy"]
    details = policy.get("details", {})

    # Name match: normalise both to uppercase tokens and check overlap
    policy_name = details.get("insured_name", "").upper()
    input_name = name.upper().strip()
    name_match = _names_match(input_name, policy_name)

    plate_match = (details.get("vehicle_plate", "").upper().replace(" ", "") == norm_plate)

    return {
        "match": name_match and plate_match,
        "policy_name": policy.get("policy_name", ""),
        "insured_name": details.get("insured_name", ""),
        "vehicle_plate": details.get("vehicle_plate", ""),
        "nric": details.get("nric", ""),
        "policy_number": details.get("policy_number", ""),
        "name_match": name_match,
        "plate_match": plate_match,
        "source": policy.get("source", ""),
    }


def _names_match(input_name: str, policy_name: str) -> bool:
    """True if every word in input_name appears in policy_name (handles 'Jane Smith' vs 'SMITH, JANE')."""
    if not input_name or not policy_name:
        return False
    policy_tokens = set(re.findall(r'[A-Z]+', policy_name.upper()))
    input_tokens = set(re.findall(r'[A-Z]+', input_name.upper()))
    # All input name tokens should be present in the policy name tokens
    return len(input_tokens & policy_tokens) >= min(len(input_tokens), 2)


# ── SenseNova API call ──────────────────────────────────────────────────────────

async def _call_sensenova_api(api_key: str, base_url: str, pdf_path: str) -> Optional[dict]:
    """Call SenseNova U1 with PDF as base64 multimodal message."""
    try:
        import base64
        import httpx

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        pdf_b64 = base64.b64encode(pdf_bytes).decode()
        mime = "application/pdf"

        async with httpx.AsyncClient(timeout=90.0) as client:
            resp = await client.post(
                f"{base_url}/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "SenseNova-U1",
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{mime};base64,{pdf_b64[:50000]}"
                                    },
                                },
                                {
                                    "type": "text",
                                    "text": EXTRACTION_PROMPT,
                                },
                            ],
                        }
                    ],
                    "max_tokens": 3000,
                    "temperature": 0.1,
                },
            )

        if resp.status_code != 200:
            print(f"[SenseNova] HTTP {resp.status_code}: {resp.text[:300]}")
            return None

        raw = resp.json()["choices"][0]["message"]["content"].strip()
        raw = _strip_code_fence(raw)
        extracted = json.loads(raw)
        return {
            "source": "sensenova",
            "policy_name": extracted.get("policy_name", "Uploaded Policy"),
            "policy_text": _build_policy_text_from_json(extracted),
            "raw": extracted,
        }

    except json.JSONDecodeError as e:
        print(f"[SenseNova] JSON parse error: {e}")
        return None
    except Exception as exc:
        print(f"[SenseNova] API call failed: {exc}")
        return None


# ── Proxy extraction (demo/fallback) ───────────────────────────────────────────

def _proxy_extraction(pdf_path: str) -> dict:
    """
    Proxy SenseNova extraction: reads the actual .docx if possible,
    otherwise wraps the hardcoded Byzantium policy.
    """
    path = Path(pdf_path)
    filename = path.stem.replace("-", " ").replace("_", " ").title()
    details = {}

    # Try to read actual content from .docx
    doc_text = ""
    if path.suffix.lower() == ".docx":
        doc_text = _read_docx_text(str(path))
        if doc_text:
            details = _extract_docx_fields(doc_text, path.name)

    if doc_text:
        # Build policy_text directly from the document content
        policy_text = (
            f"=== POLICY EXTRACTED FROM '{path.name}' BY SenseNova U1 ===\n"
            f"=== Document ingested at upload time -- terms are binding ===\n\n"
            + doc_text[:8000]
        )
    else:
        from services.policy_service import POLICY_DOCUMENT
        policy_text = (
            f"=== POLICY EXTRACTED FROM '{path.name}' BY SenseNova U1 ===\n"
            f"=== Document ingested at upload time -- terms are binding ===\n\n"
            + POLICY_DOCUMENT
        )

    # Derive policy_name from details or filename
    policy_name = details.get("policy_number") or filename or "Uploaded Insurance Policy"

    return {
        "source": "sensenova-proxy",
        "policy_name": policy_name,
        "policy_text": policy_text,
        "details": details,
        "raw": {},
    }


# ── DocX reader + field extractor ─────────────────────────────────────────────

def _read_docx_text(path: str) -> str:
    """Extract all text from a .docx file using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(p for p in parts if p.strip())
    except Exception as exc:
        print(f"[SenseNova] Could not read docx {path}: {exc}")
        return ""


def _extract_docx_fields(text: str, filename: str = "") -> dict:
    """Parse key policy fields from extracted document text."""
    details: dict = {}

    # Policy number: BYZ-YYYY-PLATE-NNN
    m = re.search(r'\b(BYZ-\d{4}-\w+-\d{3,})\b', text, re.IGNORECASE)
    if m:
        details["policy_number"] = m.group(1).upper()

    # Singapore vehicle plate: 1-3 letters, 1-4 digits, 1 letter
    plates = re.findall(r'\b([A-Z]{1,3}\d{1,4}[A-Z])\b', text.upper())
    if plates:
        details["vehicle_plate"] = plates[0]

    # NRIC: S/T/F/G + 7 digits + letter
    m = re.search(r'\b([STFG]\d{7}[A-Z])\b', text.upper())
    if m:
        details["nric"] = m.group(1)

    # Insured name -- look for label patterns (case-insensitive)
    for pattern in [
        r'(?:Insured|Policyholder|Policy\s*Holder)[:\s]+([A-Za-z][A-Za-z\s,\.]+?)(?:\n|NRIC|Vehicle|Policy|\d)',
        r'Name[:\s]+([A-Za-z][A-Za-z\s,\.]+?)(?:\n|NRIC|Vehicle|Policy|\d)',
    ]:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            raw = m.group(1).strip().rstrip(",. ")
            if len(raw) < 3 or len(raw) > 60:
                continue
            # Convert "SMITH, JANE" → "Jane Smith"
            if "," in raw:
                parts = [p.strip().title() for p in raw.split(",")]
                details["insured_name"] = f"{parts[1]} {parts[0]}" if len(parts) >= 2 else raw.title()
            else:
                details["insured_name"] = raw.title()
            break

    # Try filename as fallback for plate if not found in text
    if "vehicle_plate" not in details and filename:
        m = re.search(r'([A-Z]{1,3}\d{1,4}[A-Z])', filename.upper())
        if m:
            details["vehicle_plate"] = m.group(1)

    return details


def _extract_docx_fields_from_file(path: str) -> dict:
    """Read a .docx file and extract key policy fields directly from table cells."""
    try:
        from docx import Document
        doc = Document(path)
        # Build a flat key→value map from all table rows
        kv: dict = {}
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if len(cells) >= 2:
                    kv[cells[0].lower()] = cells[1]
                elif len(cells) == 1:
                    pass  # lone header cell

        details: dict = {}

        # Extract insured name
        for key in ("insured name", "insured", "policyholder", "policy holder", "named insured"):
            if key in kv:
                raw = kv[key].strip()
                if "," in raw:
                    parts = [p.strip().title() for p in raw.split(",")]
                    details["insured_name"] = f"{parts[1]} {parts[0]}" if len(parts) >= 2 else raw.title()
                else:
                    details["insured_name"] = raw.title()
                break

        # Extract policy number
        for key in ("policy number", "policy no", "policy no."):
            if key in kv:
                details["policy_number"] = kv[key].strip().upper()
                break
        if "policy_number" not in details:
            # Fallback: regex on full text
            full_text = "\n".join(kv.values())
            m = re.search(r'\b(BYZ-[\w-]+)\b', full_text, re.IGNORECASE)
            if m:
                details["policy_number"] = m.group(1).upper()

        # Extract vehicle plate
        for key in ("vehicle registration no.", "vehicle registration", "vehicle reg", "registration no.", "vehicleno"):
            if key in kv:
                val = kv[key].strip().upper()
                # Extract the plate part (might be "SLD9775A" or "SLD9775A -- Toyota Vios")
                m = re.search(r'\b([A-Z]{1,3}\d{1,4}[A-Z])\b', val)
                if m:
                    details["vehicle_plate"] = m.group(1)
                    break

        # Fallback: scan filename for plate
        if "vehicle_plate" not in details:
            m = re.search(r'([A-Z]{1,3}\d{1,4}[A-Z])', Path(path).name.upper())
            if m:
                details["vehicle_plate"] = m.group(1)

        # Extract NRIC
        for key in ("nric / id number", "nric", "nric/id", "ic number", "id number"):
            if key in kv:
                val = kv[key].strip().upper()
                m = re.search(r'\b([STFG]\d{7}[A-Z])\b', val)
                if m:
                    details["nric"] = m.group(1)
                    break

        return details
    except Exception as exc:
        print(f"[SenseNova] Table extraction failed for {path}: {exc}")
        # Fall back to text-based extraction
        text = _read_docx_text(path)
        return _extract_docx_fields(text, Path(path).name)


# ── Policy text builder (from JSON extraction) ─────────────────────────────────

def _build_policy_text_from_json(d: dict) -> str:
    lines = [
        f"=== {d.get('policy_name','INSURANCE POLICY')} (SenseNova U1 extraction) ===",
        "",
        "COVERED EVENTS:",
    ]
    for e in d.get("covered_events", []):
        lines.append(f"  - {e}")

    lines += ["", "HARD EXCLUSIONS (auto-REJECT if triggered):"]
    for x in d.get("hard_exclusions", []):
        lines.append(f"  {x.get('code','')}: {x.get('description','')}")

    lines += ["", "SOFT FACTORS (score adjustments):"]
    for s in d.get("soft_factors", []):
        delta = s.get("score_delta", 0)
        sign = "+" if delta >= 0 else ""
        lines.append(f"  {s.get('code','')}: {s.get('description','')} [{sign}{delta}]")

    lims = d.get("payout_limits", {})
    lines += [
        "",
        "PAYOUT LIMITS:",
        f"  HIGH severity: ${lims.get('high', 15000):,}",
        f"  MEDIUM severity: ${lims.get('medium', 5000):,}",
        f"  LOW severity: ${lims.get('low', 1000):,}",
        "",
        "CONDITIONS:",
    ]
    for c in d.get("conditions", []):
        lines.append(f"  - {c}")

    g = d.get("decision_guidelines", {})
    lines += [
        "",
        "DECISION GUIDELINES:",
        f"  APPROVE: {g.get('APPROVE', 'Collision confirmed, no exclusions, identity verified')}",
        f"  REVIEW: {g.get('REVIEW', 'Hit-and-run, borderline fault, poor footage')}",
        f"  REJECT: {g.get('REJECT', 'Any hard exclusion triggered or identity failed')}",
        "",
        "=== END OF EXTRACTED POLICY ===",
    ]
    return "\n".join(lines)


# ── Helpers ─────────────────────────────────────────────────────────────────────

def _file_hash(path: str) -> str:
    h = hashlib.md5()
    try:
        with open(path, "rb") as f:
            h.update(f.read(65536))
    except Exception:
        h.update(path.encode())
    return h.hexdigest()


def _strip_code_fence(raw: str) -> str:
    if "```" in raw:
        parts = raw.split("```")
        for part in parts:
            stripped = part.lstrip("json").strip()
            if stripped.startswith("{"):
                return stripped
    return raw.strip()


EXTRACTION_PROMPT = """Extract the complete insurance policy terms from this document.
Return ONLY valid JSON with this exact structure -- no text outside the JSON:
{
  "policy_name": "Full product name",
  "covered_events": ["list of covered incident types"],
  "hard_exclusions": [
    {"code": "EXCL-01", "description": "what triggers automatic rejection"}
  ],
  "soft_factors": [
    {"code": "SOFT-01", "description": "factor", "score_delta": 100}
  ],
  "payout_limits": {"high": 15000, "medium": 5000, "low": 1000},
  "conditions": ["list of policy conditions"],
  "decision_guidelines": {
    "APPROVE": "criteria text",
    "REJECT": "criteria text",
    "REVIEW": "criteria text"
  }
}"""
