"""Generate demo insurance policy .docx for Jane Smith / SLD9775A"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date, timedelta
import os

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────────
def heading(doc, text, level=1, color=RGBColor(0x1A, 0x56, 0xDB)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def para(doc, text, bold=False, size=10, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def rule_row(table, code, description, value=""):
    row = table.add_row()
    row.cells[0].text = code
    row.cells[1].text = description
    if value:
        row.cells[2].text = value
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    if value:
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)

def add_table_header(table, *cols):
    hdr = table.rows[0]
    for i, col in enumerate(cols):
        cell = hdr.cells[i]
        cell.text = col
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Set cell background to dark blue
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A56DB')
        tcPr.append(shd)

# ── Cover Header ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BYZANTIUM INSURANCE PTE. LTD.")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Motor Insurance Policy Certificate")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 72)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run.font.size = Pt(9)

doc.add_paragraph()

# ── Policy Summary Table ──────────────────────────────────────────────────────────
heading(doc, "POLICY DETAILS", level=2)

summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = "Table Grid"
summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT

issue_date  = date(2024, 1, 15)
expiry_date = date(2025, 1, 14)

rows_data = [
    ("Policy Number",        "BYZ-2024-SLD9775A-001"),
    ("Certificate Number",   "CERT-SLD9775A-2024"),
    ("Policy Type",          "Comprehensive Motor Insurance"),
    ("Insured Name",         "SMITH, JANE"),
    ("NRIC / ID Number",     "S9812381D"),
    ("Date of Birth",        "12 March 1998"),
    ("Address",              "Blk 45 Ang Mo Kio Ave 3, #08-22, Singapore 560045"),
    ("Contact Number",       "+65 9123 4567"),
    ("Email",                "jane.smith@email.com"),
    ("Policy Issued",        issue_date.strftime("%d %B %Y")),
    ("Policy Effective",     issue_date.strftime("%d %B %Y")),
    ("Policy Expiry",        expiry_date.strftime("%d %B %Y")),
    ("Annual Premium",       "SGD 1,420.00 (inclusive of GST)"),
    ("Policy Status",        "ACTIVE"),
]

for label, value in rows_data:
    row = summary_table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(9)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(9)
        if label == "Policy Status":
            run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            run.font.bold = True

summary_table.rows[0].cells[0].text = ""
summary_table.rows[0].cells[1].text = ""

doc.add_paragraph()

# ── Insured Vehicle ───────────────────────────────────────────────────────────────
heading(doc, "INSURED VEHICLE", level=2)

vehicle_table = doc.add_table(rows=1, cols=2)
vehicle_table.style = "Table Grid"
vehicle_table.alignment = WD_TABLE_ALIGNMENT.LEFT

vehicle_rows = [
    ("Vehicle Registration No.",  "SLD9775A"),
    ("Make / Model",              "Toyota Vios 1.5G"),
    ("Year of Manufacture",       "2019"),
    ("Engine Capacity",           "1,497 cc"),
    ("Vehicle Type",              "Private Saloon Car"),
    ("Engine No.",                "2NR-FE-19X04821"),
    ("Chassis No.",               "MHFEX59G4K0012345"),
    ("Colour",                    "Pearl White"),
    ("Agreed Value / Sum Insured","SGD 52,000"),
    ("NCD (No-Claim Discount)",   "30%"),
    ("Dashcam Installed",         "Yes — Front & Rear (Blackvue DR750X)"),
]

for label, value in vehicle_rows:
    row = vehicle_table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    for run in row.cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(9)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(9)
        if label == "Vehicle Registration No.":
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

vehicle_table.rows[0].cells[0].text = ""
vehicle_table.rows[0].cells[1].text = ""

doc.add_paragraph()

# ── Named Drivers ─────────────────────────────────────────────────────────────────
heading(doc, "NAMED DRIVERS", level=2)

driver_table = doc.add_table(rows=1, cols=4)
driver_table.style = "Table Grid"
add_table_header(driver_table, "Driver Name", "NRIC", "Driving Licence", "Relationship")

driver_data = [
    ("SMITH, JANE (Principal)",  "S9812381D", "Class 3 — Valid",  "Named Insured"),
]
for row_data in driver_data:
    row = driver_table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        run = row.cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(9)

doc.add_paragraph()

# ── Section 1: Covered Events ─────────────────────────────────────────────────────
heading(doc, "SECTION 1 — COVERED EVENTS", level=2)

covered = [
    "Accidental loss or damage to the insured vehicle caused by collision or overturning",
    "Third-party bodily injury or death arising from use of insured vehicle",
    "Third-party property damage arising from use of insured vehicle",
    "Rear-end collision — claimant stationary or in slow-moving traffic",
    "Side-impact (T-bone) or intersection collision where claimant was not the aggressor",
    "Head-on collision where claimant was NOT the primary at-fault driver",
    "Multi-vehicle pile-up involvement",
    "Weather-induced collision (ice, flooding, low-visibility fog, rain)",
    "Hit-and-run by untraced third party (subject to REVIEW and police report)",
    "Fire and theft of the insured vehicle",
    "Flood and storm damage to insured vehicle",
]

for item in covered:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(item)
    run.font.size = Pt(9)

doc.add_paragraph()

# ── Section 2: Hard Exclusions ────────────────────────────────────────────────────
heading(doc, "SECTION 2 — HARD EXCLUSIONS (Automatic Rejection — No Exceptions)", level=2)
para(doc, "Any claim triggering one or more of the following exclusions will be automatically REJECTED "
     "without further assessment.", size=9, color=RGBColor(0x99, 0x00, 0x00))
doc.add_paragraph()

excl_table = doc.add_table(rows=1, cols=3)
excl_table.style = "Table Grid"
add_table_header(excl_table, "Code", "Exclusion", "Consequence")

exclusions = [
    ("EXCL-01", "Claimant was the at-fault driver in a head-on or primary-aggressor collision",               "Auto-REJECT"),
    ("EXCL-02", "Incident occurred in a private car park at high speed (>30 km/h)",                          "Auto-REJECT"),
    ("EXCL-03", "Evidence of distracted driving — phone use, reckless behaviour visible in dashcam footage", "Auto-REJECT"),
    ("EXCL-04", "Vehicle type in footage does not match insured vehicle class (SLD9775A — Private Saloon)",   "Auto-REJECT"),
    ("EXCL-05", "Single-vehicle accident with no identifiable external cause (e.g. self-inflicted, barrier)", "Auto-REJECT"),
    ("EXCL-06", "Dashcam footage shows pre-existing damage inconsistent with the claimed incident",           "Auto-REJECT"),
    ("EXCL-07", "Claim filed more than 30 calendar days after the date of incident",                         "Auto-REJECT"),
    ("EXCL-08", "Claimant identity cannot be verified via approved methods (Terminal 3 / SingPass MyInfo)",   "Auto-REJECT"),
]

for code, desc, cons in exclusions:
    row = excl_table.add_row()
    row.cells[0].text = code
    row.cells[1].text = desc
    row.cells[2].text = cons
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    row.cells[2].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ── Section 3: Soft Factors ───────────────────────────────────────────────────────
heading(doc, "SECTION 3 — ASSESSMENT FACTORS (Trust Score Adjustments)", level=2)
para(doc, "These factors are applied by the AI claims officer to adjust the trust score. "
     "Score range: 0–1000. APPROVE ≥ 700 | REVIEW 400–699 | REJECT < 400.", size=9)
doc.add_paragraph()

soft_table = doc.add_table(rows=1, cols=3)
soft_table.style = "Table Grid"
add_table_header(soft_table, "Code", "Factor Description", "Score Delta")

soft_factors = [
    ("SOFT-01", "Claimant vehicle was stationary at time of impact — strong indicator of non-fault",    "+100"),
    ("SOFT-02", "Adverse weather or significantly reduced visibility at time of incident",                "+50"),
    ("SOFT-03", "Three or more vehicles visible in footage — corroborating evidence",                    "+50"),
    ("SOFT-04", "Dashcam clearly shows fault of other party — high evidence quality",                    "+100"),
    ("SOFT-05", "Severity LOW with no visible structural damage — reduced payout justification",         "–75"),
    ("SOFT-06", "Night-time incident with poor dashcam footage quality — reduced evidence confidence",   "–50"),
    ("SOFT-07", "Hit-and-run — offending vehicle fled, no second party identifiable (REVIEW required)",  "–100"),
    ("SOFT-08", "Single vehicle incident without clear external cause — suspicious",                      "–150"),
]

for code, desc, delta in soft_factors:
    row = soft_table.add_row()
    row.cells[0].text = code
    row.cells[1].text = desc
    row.cells[2].text = delta
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[2].paragraphs[0].runs[0].font.bold = True
    color = RGBColor(0x05, 0x96, 0x69) if delta.startswith("+") else RGBColor(0x99, 0x00, 0x00)
    row.cells[2].paragraphs[0].runs[0].font.color.rgb = color

doc.add_paragraph()

# ── Section 4: Payout Limits ──────────────────────────────────────────────────────
heading(doc, "SECTION 4 — PAYOUT LIMITS", level=2)

payout_table = doc.add_table(rows=1, cols=3)
payout_table.style = "Table Grid"
add_table_header(payout_table, "Severity", "Maximum Single Claim Payout", "Conditions")

payout_data = [
    ("HIGH",   "SGD 15,000", "Major structural damage, airbag deployment, write-off"),
    ("MEDIUM", "SGD 5,000",  "Moderate damage, panel deformation, repairable"),
    ("LOW",    "SGD 1,000",  "Minor damage, scratches, dents — no structural impact"),
    ("NONE",   "SGD 0",      "No collision confirmed — claim rejected or not applicable"),
]

for sev, amount, cond in payout_data:
    row = payout_table.add_row()
    row.cells[0].text = sev
    row.cells[1].text = amount
    row.cells[2].text = cond
    color = (RGBColor(0x99, 0x00, 0x00) if sev == "HIGH"
             else RGBColor(0xD9, 0x77, 0x06) if sev == "MEDIUM"
             else RGBColor(0x05, 0x96, 0x69) if sev == "LOW"
             else RGBColor(0x66, 0x66, 0x66))
    for cell in row.cells:
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = color
    row.cells[1].paragraphs[0].runs[0].font.bold = True

p = doc.add_paragraph()
run = p.add_run("Annual claim limit: SGD 25,000 per policy. Multiple claims within 30 days are subject to fraud review.")
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

# ── Section 5: Coverage Conditions ────────────────────────────────────────────────
heading(doc, "SECTION 5 — COVERAGE CONDITIONS", level=2)

conditions = [
    ("COND-01", "This policy must be active and in-force at the date and time of the incident."),
    ("COND-02", "Only the named insured (Jane Smith, S9812381D) or a listed additional driver may operate the insured vehicle."),
    ("COND-03", "All incidents must be reported to Byzantium Insurance within 30 calendar days of occurrence."),
    ("COND-04", "Dashcam footage submitted must be original, unedited, and timestamped from the incident vehicle (SLD9775A)."),
    ("COND-05", "Claimant identity must be verifiable via Terminal 3 TEE attestation or SingPass MyInfo government data."),
    ("COND-06", "A valid police report is required for all third-party claims and hit-and-run incidents."),
    ("COND-07", "The insured vehicle must be roadworthy and hold a valid LTA Certificate of Entitlement (COE)."),
]

for code, text in conditions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.0)
    run_code = p.add_run(f"{code}  ")
    run_code.font.bold = True
    run_code.font.size = Pt(9)
    run_code.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run_text = p.add_run(text)
    run_text.font.size = Pt(9)

doc.add_paragraph()

# ── Section 6: Decision Guidelines ────────────────────────────────────────────────
heading(doc, "SECTION 6 — AI CLAIMS DECISION GUIDELINES", level=2)

dec_table = doc.add_table(rows=1, cols=3)
dec_table.style = "Table Grid"
add_table_header(dec_table, "Decision", "Trust Score", "Criteria")

decisions = [
    ("APPROVE", "700 – 1000",
     "No hard exclusions triggered. Collision confirmed. Identity verified (T3 score ≥ 75). "
     "Dashcam footage of adequate quality. Positive soft factors present."),
    ("REVIEW",  "400 – 699",
     "Hit-and-run incident (SOFT-07). Borderline fault determination. Poor footage quality. "
     "Fraud signals detected (cross-claim checks). Identity borderline (T3 score 75–79)."),
    ("REJECT",  "0 – 399",
     "Any hard exclusion triggered (EXCL-01 to EXCL-08). No collision evidence. "
     "Identity unverified. Suspended driving licence. Duplicate claim within 30 days."),
]

for dec, score, criteria in decisions:
    row = dec_table.add_row()
    row.cells[0].text = dec
    row.cells[1].text = score
    row.cells[2].text = criteria
    color = (RGBColor(0x05, 0x96, 0x69) if dec == "APPROVE"
             else RGBColor(0xD9, 0x77, 0x06) if dec == "REVIEW"
             else RGBColor(0x99, 0x00, 0x00))
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = color
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── Section 7: Fraud Detection ────────────────────────────────────────────────────
heading(doc, "SECTION 7 — AUTOMATED FRAUD DETECTION", level=2)
para(doc, "All claims are subject to real-time cross-claim fraud screening by the Daytona ClaimAgent.", size=9)
doc.add_paragraph()

fraud_table = doc.add_table(rows=1, cols=3)
fraud_table.style = "Table Grid"
add_table_header(fraud_table, "Flag", "Trigger Condition", "Action")

fraud_rules = [
    ("FRAUD-01", "Same policy number filed twice within 30 calendar days",                     "Auto-REJECT — hard rule"),
    ("FRAUD-02", "Second claim on same policy within 24 hours of a prior claim",               "Force REVIEW + score –200"),
    ("FRAUD-03", "Same vehicle plate fingerprint detected across multiple distinct claims",     "Force REVIEW + score –150"),
]

for flag, trigger, action in fraud_rules:
    row = fraud_table.add_row()
    row.cells[0].text = flag
    row.cells[1].text = trigger
    row.cells[2].text = action
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── Section 8: Technology Stack ───────────────────────────────────────────────────
heading(doc, "SECTION 8 — CLAIM PROCESSING PIPELINE", level=2)
para(doc, "This policy is processed through the Byzantium AutoClaims AI pipeline:", size=9)
doc.add_paragraph()

stack = [
    ("Nosana GPU",      "Decentralized GPU compute — CLIP-ViT-B/32 frame analysis, collision signal detection, footage integrity verification"),
    ("VideoDB",         "Dashcam video indexing — shot-based scene extraction, spoken word audio indexing, collision clip generation"),
    ("Terminal 3",      "TEE-based claimant identity attestation (DID: did:t3n:16d34b1887e89257702b597ef585298db72cad82)"),
    ("SingPass MyInfo", "Government-verified vehicle ownership, driving licence status, demerit points (NRIC: S9812381D)"),
    ("SenseNova U1",    "Multimodal policy document ingestion — extracts coverage terms from PDF/PPTX for Kimi reasoning"),
    ("Kimi AI",         "AI claims officer — cross-references all evidence against this policy document to make APPROVE/REJECT/REVIEW"),
    ("Daytona",         "Secure ClaimAgent sandbox — enforces hard policy exclusions, soft factor scoring, and fraud detection rules"),
]

for tool, desc in stack:
    p = doc.add_paragraph()
    run_tool = p.add_run(f"{tool}: ")
    run_tool.font.bold = True
    run_tool.font.size = Pt(9)
    run_tool.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run_desc = p.add_run(desc)
    run_desc.font.size = Pt(9)

doc.add_paragraph()

# ── Footer ────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 72)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Byzantium Insurance Pte. Ltd.  ·  UEN: 202400001Z  ·  "
    "MAS-Licensed General Insurer  ·  policies@byzantium.sg"
)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Policy generated: {date.today().strftime('%d %B %Y')}  ·  "
    "This is a digitally processed document. Verify authenticity at claims.byzantium.sg"
)
run.font.size = Pt(7)
run.font.italic = True
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Save ──────────────────────────────────────────────────────────────────────────
out_dir = r"c:\Users\aadha\OneDrive\Desktop\projects\Agent Hackathon 130626"
out_path = os.path.join(out_dir, "JaneSmith_SLD9775A_Policy.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
